"""
Export service for generating professional DOCX and PDF files from proposal content.
Produces polished, government-ready documents with tables, matrices, and proper formatting.
"""

import io
import logging
import re
from datetime import datetime
from typing import Dict, List, Optional, Tuple
from html import unescape as html_unescape

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph as RLParagraph,
    Spacer,
    Table as RLTable,
    TableStyle,
    PageBreak,
    KeepTogether,
)

logger = logging.getLogger(__name__)

# Color palette
NAVY = RGBColor(0x1B, 0x2A, 0x4A)
NAVY_HEX = "1B2A4A"
ACCENT = RGBColor(0x10, 0xB9, 0x81)
ACCENT_HEX = "10B981"
DARK_TEXT = RGBColor(0x37, 0x41, 0x51)
GRAY_TEXT = RGBColor(0x6B, 0x72, 0x80)
LIGHT_BG = "F0FDF4"
HEADER_BG = "1B2A4A"
SUB_HEADER_BG = "059669"
LIGHT_GRAY_BG = "F3F4F6"
BORDER_COLOR = RGBColor(0xD1, 0xD5, 0xDB)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def _set_cell_shading(cell, color_hex: str):
    """Set cell background color."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def _set_cell_borders(cell, color="D1D5DB", size="4"):
    """Set cell borders."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:left w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:bottom w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:right w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(tcBorders)


def _add_formatted_paragraph(doc, text: str, size=11, color=DARK_TEXT, bold=False,
                              italic=False, alignment=None, space_before=0, space_after=6,
                              font_name="Calibri"):
    """Add a paragraph with specific formatting."""
    para = doc.add_paragraph()
    if alignment:
        para.alignment = alignment
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after = Pt(space_after)

    # Handle markdown bold within text
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = para.add_run(part[2:-2])
            run.bold = True
        else:
            run = para.add_run(part)
            run.bold = bold

    for run in para.runs:
        run.font.size = Pt(size)
        run.font.color.rgb = color
        run.font.name = font_name
        if italic:
            run.italic = True
    return para


def _add_styled_table(doc, headers: List[str], rows: List[List[str]],
                       header_bg=HEADER_BG, alt_row_bg=LIGHT_GRAY_BG):
    """Add a professionally styled table."""
    if not headers or not rows:
        return None

    num_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, header_text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header_text)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = WHITE
        run.font.name = "Calibri"
        _set_cell_shading(cell, header_bg)
        _set_cell_borders(cell, NAVY_HEX)

    # Data rows
    for row_idx, row_data in enumerate(rows):
        for col_idx in range(num_cols):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell_text = row_data[col_idx] if col_idx < len(row_data) else ""
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(cell_text)
            run.font.size = Pt(10)
            run.font.color.rgb = DARK_TEXT
            run.font.name = "Calibri"
            _set_cell_borders(cell)
            if row_idx % 2 == 1:
                _set_cell_shading(cell, alt_row_bg)

    return table


def _detect_table_content(lines: List[str]) -> Optional[Tuple[List[str], List[List[str]]]]:
    """Detect if a block of lines contains table-like data (pipe-separated or consistent delimiters)."""
    # Check for pipe-separated tables: | Header1 | Header2 |
    pipe_lines = [l for l in lines if '|' in l and l.count('|') >= 2]
    if len(pipe_lines) >= 2:
        headers = []
        rows = []
        for i, line in enumerate(pipe_lines):
            cells = [c.strip() for c in line.strip('|').split('|')]
            cells = [c for c in cells if c and not re.match(r'^[-:]+$', c)]
            if not cells:
                continue
            if not headers:
                headers = cells
            else:
                rows.append(cells)
        if headers and rows:
            return headers, rows

    # Check for colon-separated key-value pairs (3+ consecutive)
    kv_lines = []
    for line in lines:
        if ':' in line and not line.startswith('#'):
            parts = line.split(':', 1)
            if len(parts) == 2 and len(parts[0].strip()) > 1 and len(parts[1].strip()) > 1:
                kv_lines.append((parts[0].strip(), parts[1].strip()))
    if len(kv_lines) >= 3:
        return ["Item", "Details"], [[k, v] for k, v in kv_lines]

    return None


def _detect_risk_register(text: str) -> Optional[Tuple[List[str], List[List[str]]]]:
    """Detect risk register format: Risk: X | Probability: Y | Impact: Z"""
    pattern = r'(?:Risk|Issue)\s*:\s*(.+?)\s*\|\s*(?:Probability|Likelihood)\s*:\s*(\w+)\s*\|\s*Impact\s*:\s*(\w+)\s*\|\s*Mitigation\s*:\s*(.+?)(?:\s*\|\s*Contingency\s*:\s*(.+?))?(?:\n|$)'
    matches = re.findall(pattern, text, re.IGNORECASE)
    if matches:
        headers = ["Risk", "Probability", "Impact", "Mitigation"]
        if any(m[4] for m in matches):
            headers.append("Contingency")
        rows = []
        for m in matches:
            row = [m[0].strip(), m[1].strip(), m[2].strip(), m[3].strip()]
            if len(headers) == 5:
                row.append(m[4].strip() if m[4] else "")
            rows.append(row)
        return headers, rows
    return None


def generate_docx(proposal: Dict) -> io.BytesIO:
    """Generate a professional Word document from proposal sections."""
    try:
        doc = Document()

        # --- Document styles ---
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(6)

        proposal_title = proposal.get("proposal_title", "Government Proposal")
        vendor_name = proposal.get("vendor_name", "")
        sections = proposal.get("sections", {})

        # ============================================================
        # COVER PAGE
        # ============================================================
        for _ in range(3):
            doc.add_paragraph("")

        # Accent line
        line_table = doc.add_table(rows=1, cols=1)
        line_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = line_table.rows[0].cells[0]
        cell.text = ""
        _set_cell_shading(cell, ACCENT_HEX)
        cell.width = Inches(5)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(" ")
        run.font.size = Pt(2)

        doc.add_paragraph("")

        # Main title
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_para.paragraph_format.space_after = Pt(8)
        title_run = title_para.add_run(proposal_title.upper())
        title_run.bold = True
        title_run.font.size = Pt(28)
        title_run.font.color.rgb = NAVY
        title_run.font.name = "Calibri"

        # Subtitle
        subtitle_para = doc.add_paragraph()
        subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_run = subtitle_para.add_run("Government Contract Proposal")
        sub_run.font.size = Pt(14)
        sub_run.font.color.rgb = ACCENT
        sub_run.font.name = "Calibri"

        if vendor_name:
            doc.add_paragraph("")
            _add_formatted_paragraph(
                doc, f"Prepared by: {vendor_name}",
                size=14, color=NAVY, bold=True,
                alignment=WD_ALIGN_PARAGRAPH.CENTER
            )

        # Date
        _add_formatted_paragraph(
            doc, f"Date: {datetime.now().strftime('%B %d, %Y')}",
            size=12, color=GRAY_TEXT,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            space_before=12
        )

        # Spacer
        for _ in range(4):
            doc.add_paragraph("")

        # Confidential notice
        _add_formatted_paragraph(
            doc, "CONFIDENTIAL - FOR GOVERNMENT USE ONLY",
            size=10, color=RGBColor(0x99, 0x00, 0x00), italic=True,
            alignment=WD_ALIGN_PARAGRAPH.CENTER
        )

        doc.add_page_break()

        # ============================================================
        # TABLE OF CONTENTS
        # ============================================================
        toc_heading = doc.add_heading("Table of Contents", level=1)
        for run in toc_heading.runs:
            run.font.color.rgb = NAVY
            run.font.name = "Calibri"

        doc.add_paragraph("")

        # TOC as a clean table
        toc_table = doc.add_table(rows=len(sections), cols=2)
        toc_table.alignment = WD_TABLE_ALIGNMENT.LEFT
        for idx, (section_key, section_data) in enumerate(sections.items()):
            section_title = section_data.get("title", section_key.replace("_", " ").title())
            # Number cell
            num_cell = toc_table.rows[idx].cells[0]
            num_cell.text = ""
            p = num_cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(f"{idx + 1}.")
            run.font.size = Pt(11)
            run.font.color.rgb = ACCENT
            run.font.bold = True
            run.font.name = "Calibri"
            num_cell.width = Inches(0.5)

            # Title cell
            title_cell = toc_table.rows[idx].cells[1]
            title_cell.text = ""
            p = title_cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(section_title)
            run.font.size = Pt(11)
            run.font.color.rgb = DARK_TEXT
            run.font.name = "Calibri"

            # Alternating row shading
            if idx % 2 == 0:
                _set_cell_shading(num_cell, LIGHT_GRAY_BG)
                _set_cell_shading(title_cell, LIGHT_GRAY_BG)

        doc.add_page_break()

        # ============================================================
        # CONTENT SECTIONS
        # ============================================================
        for idx, (section_key, section_data) in enumerate(sections.items(), start=1):
            section_title = section_data.get("title", section_key.replace("_", " ").title())
            content = section_data.get("content", "")

            if idx > 1:
                doc.add_page_break()

            # Section number badge + title
            heading_para = doc.add_paragraph()
            heading_para.paragraph_format.space_before = Pt(0)
            heading_para.paragraph_format.space_after = Pt(2)

            num_run = heading_para.add_run(f"  {idx}  ")
            num_run.bold = True
            num_run.font.size = Pt(12)
            num_run.font.color.rgb = WHITE
            num_run.font.name = "Calibri"
            # We can't easily add a background to a run in python-docx,
            # so we use a different visual approach

            # Section heading with accent underline
            section_heading = doc.add_heading(f"{idx}. {section_title}", level=1)
            for run in section_heading.runs:
                run.font.color.rgb = NAVY
                run.font.name = "Calibri"
            # Remove the auto-generated empty paragraph
            heading_para._element.getparent().remove(heading_para._element)

            # Accent underline
            line_table = doc.add_table(rows=1, cols=1)
            line_table.alignment = WD_TABLE_ALIGNMENT.LEFT
            cell = line_table.rows[0].cells[0]
            cell.text = ""
            _set_cell_shading(cell, ACCENT_HEX)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(" ")
            run.font.size = Pt(1)

            doc.add_paragraph("")

            # Check for risk register format in the full content
            risk_data = _detect_risk_register(content)
            if risk_data and section_key in ('risk_mitigation', 'risk_register'):
                # Render narrative text before the table
                pre_table_lines = []
                for line in content.split('\n'):
                    if re.search(r'(?:Risk|Issue)\s*:', line, re.IGNORECASE) and '|' in line:
                        break
                    pre_table_lines.append(line)
                if pre_table_lines:
                    _render_content_lines(doc, '\n'.join(pre_table_lines))
                _add_styled_table(doc, risk_data[0], risk_data[1])
                continue

            # Parse and render content
            _render_content_lines(doc, content, section_key)

        # ============================================================
        # FOOTER - Generated notice
        # ============================================================
        doc.add_paragraph("")
        _add_formatted_paragraph(
            doc,
            f"Generated by GovProposal AI on {datetime.now().strftime('%B %d, %Y')}",
            size=8, color=GRAY_TEXT, italic=True,
            alignment=WD_ALIGN_PARAGRAPH.CENTER
        )

        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    except Exception as exc:
        logger.error("Failed to generate DOCX: %s", exc)
        raise RuntimeError(f"DOCX generation failed: {exc}")


def _render_content_lines(doc, content: str, section_key: str = ""):
    """Render content with smart detection of tables, lists, and structured data."""
    parsed = _html_to_paragraphs(content)

    # Collect consecutive lines to detect table blocks
    i = 0
    while i < len(parsed):
        para = parsed[i]
        ptype = para['type']
        ptext = para['text']

        if ptype == 'spacer':
            doc.add_paragraph("")
            i += 1
            continue

        # Check if this starts a table-like block (pipe-separated lines)
        if '|' in ptext and ptext.count('|') >= 2:
            table_lines = [ptext]
            j = i + 1
            while j < len(parsed) and ('|' in parsed[j]['text'] and parsed[j]['text'].count('|') >= 2 or re.match(r'^[-|:]+$', parsed[j]['text'].strip())):
                table_lines.append(parsed[j]['text'])
                j += 1
            if len(table_lines) >= 2:
                table_data = _detect_table_content(table_lines)
                if table_data:
                    _add_styled_table(doc, table_data[0], table_data[1])
                    doc.add_paragraph("")
                    i = j
                    continue

        # Regular content rendering
        if ptype in ('h1', 'h2'):
            heading = doc.add_heading(_strip_markdown_bold(ptext), level=2)
            for run in heading.runs:
                run.font.color.rgb = RGBColor(0x05, 0x96, 0x69)
                run.font.name = "Calibri"
        elif ptype == 'h3':
            heading = doc.add_heading(_strip_markdown_bold(ptext), level=3)
            for run in heading.runs:
                run.font.color.rgb = DARK_TEXT
                run.font.name = "Calibri"
        elif ptype == 'bullet':
            clean = _strip_markdown_bold(ptext)
            if clean:
                bp = doc.add_paragraph(style="List Bullet")
                # Handle bold parts within bullet
                parts = re.split(r'(\*\*.*?\*\*)', ptext)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = bp.add_run(part[2:-2])
                        run.bold = True
                    else:
                        run = bp.add_run(part)
                    run.font.size = Pt(11)
                    run.font.color.rgb = DARK_TEXT
                    run.font.name = "Calibri"
        elif ptype == 'numbered':
            clean = _strip_markdown_bold(ptext)
            if clean:
                np = doc.add_paragraph(style="List Number")
                run = np.add_run(clean)
                run.font.size = Pt(11)
                run.font.color.rgb = DARK_TEXT
                run.font.name = "Calibri"
        else:
            clean = ptext.strip()
            if clean:
                _add_formatted_paragraph(doc, clean, size=11, color=DARK_TEXT)

        i += 1


# ============================================================
# PDF Export
# ============================================================

def generate_pdf(proposal: Dict) -> io.BytesIO:
    """Generate a professional PDF document from proposal sections."""
    try:
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=letter,
            rightMargin=0.75 * inch,
            leftMargin=0.75 * inch,
            topMargin=0.75 * inch,
            bottomMargin=0.75 * inch,
        )

        styles = getSampleStyleSheet()

        styles.add(ParagraphStyle(
            name="ProposalTitle", parent=styles["Title"],
            fontSize=26, textColor=colors.HexColor("#1B2A4A"),
            spaceAfter=12, alignment=1, fontName="Helvetica-Bold",
        ))
        styles.add(ParagraphStyle(
            name="ProposalSubtitle", parent=styles["Normal"],
            fontSize=14, textColor=colors.HexColor("#10B981"),
            spaceAfter=8, alignment=1, fontName="Helvetica",
        ))
        styles.add(ParagraphStyle(
            name="SectionHeading", parent=styles["Heading1"],
            fontSize=16, textColor=colors.HexColor("#1B2A4A"),
            spaceBefore=20, spaceAfter=10, fontName="Helvetica-Bold",
        ))
        styles.add(ParagraphStyle(
            name="SubHeading", parent=styles["Heading2"],
            fontSize=13, textColor=colors.HexColor("#059669"),
            spaceBefore=12, spaceAfter=6, fontName="Helvetica-Bold",
        ))
        styles.add(ParagraphStyle(
            name="ProposalBody", parent=styles["Normal"],
            fontSize=10.5, leading=14, spaceAfter=6,
            fontName="Helvetica", textColor=colors.HexColor("#374151"),
        ))
        styles.add(ParagraphStyle(
            name="ProposalBullet", parent=styles["Normal"],
            fontSize=10.5, leading=14, spaceAfter=4,
            leftIndent=20, bulletIndent=10,
            fontName="Helvetica", textColor=colors.HexColor("#374151"),
        ))
        styles.add(ParagraphStyle(
            name="Confidential", parent=styles["Normal"],
            fontSize=9, textColor=colors.HexColor("#990000"),
            alignment=1, fontName="Helvetica-Oblique",
        ))
        styles.add(ParagraphStyle(
            name="TableHeader", parent=styles["Normal"],
            fontSize=9, textColor=colors.white,
            fontName="Helvetica-Bold", alignment=1,
        ))
        styles.add(ParagraphStyle(
            name="TableCell", parent=styles["Normal"],
            fontSize=9, textColor=colors.HexColor("#374151"),
            fontName="Helvetica", leading=12,
        ))

        story = []
        proposal_title = proposal.get("proposal_title", "Government Proposal")
        vendor_name = proposal.get("vendor_name", "")
        sections = proposal.get("sections", {})

        # --- Cover Page ---
        story.append(Spacer(1, 1.5 * inch))

        # Accent line
        accent_line = RLTable([[""]], colWidths=[5 * inch], rowHeights=[3])
        accent_line.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#10B981")),
            ("ALIGN", (0, 0), (-1, -1), "CENTER"),
        ]))
        story.append(accent_line)
        story.append(Spacer(1, 0.5 * inch))

        story.append(RLParagraph(_escape_xml(proposal_title.upper()), styles["ProposalTitle"]))
        story.append(Spacer(1, 0.15 * inch))
        story.append(RLParagraph("Government Contract Proposal", styles["ProposalSubtitle"]))

        if vendor_name:
            story.append(Spacer(1, 0.3 * inch))
            vendor_style = ParagraphStyle(
                "VendorName", parent=styles["Normal"],
                fontSize=14, textColor=colors.HexColor("#1B2A4A"),
                alignment=1, fontName="Helvetica-Bold",
            )
            story.append(RLParagraph(f"Prepared by: {_escape_xml(vendor_name)}", vendor_style))

        date_style = ParagraphStyle(
            "DateLine", parent=styles["Normal"],
            fontSize=12, textColor=colors.HexColor("#6B7280"),
            alignment=1, fontName="Helvetica",
        )
        story.append(Spacer(1, 0.15 * inch))
        story.append(RLParagraph(f"Date: {datetime.now().strftime('%B %d, %Y')}", date_style))

        story.append(Spacer(1, 2 * inch))
        story.append(RLParagraph("CONFIDENTIAL - FOR GOVERNMENT USE ONLY", styles["Confidential"]))
        story.append(PageBreak())

        # --- Table of Contents ---
        story.append(RLParagraph("Table of Contents", styles["SectionHeading"]))

        # Accent line under TOC heading
        toc_line = RLTable([[""]], colWidths=[6.5 * inch], rowHeights=[2])
        toc_line.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#10B981")),
        ]))
        story.append(toc_line)
        story.append(Spacer(1, 0.2 * inch))

        # TOC as styled table
        toc_data = []
        for idx, (section_key, section_data) in enumerate(sections.items(), start=1):
            section_title = section_data.get("title", section_key.replace("_", " ").title())
            toc_data.append([
                RLParagraph(f"<b>{idx}.</b>", styles["TableCell"]),
                RLParagraph(_escape_xml(section_title), styles["TableCell"]),
            ])

        if toc_data:
            toc_table = RLTable(toc_data, colWidths=[0.5 * inch, 5.5 * inch])
            toc_style = [
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]
            for i in range(0, len(toc_data), 2):
                toc_style.append(("BACKGROUND", (0, i), (-1, i), colors.HexColor("#F3F4F6")))
            toc_table.setStyle(TableStyle(toc_style))
            story.append(toc_table)

        story.append(PageBreak())

        # --- Content Sections ---
        for idx, (section_key, section_data) in enumerate(sections.items(), start=1):
            section_title = section_data.get("title", section_key.replace("_", " ").title())
            content = section_data.get("content", "")

            if idx > 1:
                story.append(PageBreak())

            # Section heading
            story.append(RLParagraph(
                f"{idx}. {_escape_xml(section_title)}",
                styles["SectionHeading"],
            ))

            # Accent line under heading
            line = RLTable([[""]], colWidths=[6.5 * inch], rowHeights=[2])
            line.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#10B981")),
            ]))
            story.append(line)
            story.append(Spacer(1, 0.15 * inch))

            # Parse and render content
            parsed = _html_to_paragraphs(content)
            _render_pdf_content(story, parsed, styles, section_key)
            story.append(Spacer(1, 0.3 * inch))

        # Footer
        story.append(Spacer(1, 0.5 * inch))
        footer_style = ParagraphStyle(
            "Footer", parent=styles["Normal"],
            fontSize=8, textColor=colors.HexColor("#6B7280"),
            alignment=1, fontName="Helvetica-Oblique",
        )
        story.append(RLParagraph(
            f"Generated by GovProposal AI on {datetime.now().strftime('%B %d, %Y')}",
            footer_style
        ))

        doc.build(story)
        buffer.seek(0)
        return buffer

    except Exception as exc:
        logger.error("Failed to generate PDF: %s", exc)
        raise RuntimeError(f"PDF generation failed: {exc}")


def _render_pdf_content(story, parsed, styles, section_key=""):
    """Render parsed content for PDF with smart table detection."""
    i = 0
    while i < len(parsed):
        para = parsed[i]
        ptype = para['type']
        ptext = para['text']

        if ptype == 'spacer':
            story.append(Spacer(1, 4))
            i += 1
            continue

        # Detect table blocks
        if '|' in ptext and ptext.count('|') >= 2:
            table_lines = [ptext]
            j = i + 1
            while j < len(parsed) and ('|' in parsed[j]['text'] and parsed[j]['text'].count('|') >= 2 or re.match(r'^[-|:]+$', parsed[j]['text'].strip())):
                table_lines.append(parsed[j]['text'])
                j += 1
            if len(table_lines) >= 2:
                table_data = _detect_table_content(table_lines)
                if table_data:
                    _add_pdf_table(story, table_data[0], table_data[1], styles)
                    story.append(Spacer(1, 8))
                    i = j
                    continue

        if ptype in ('h1', 'h2'):
            story.append(RLParagraph(_escape_xml(ptext), styles["SubHeading"]))
        elif ptype == 'h3':
            story.append(RLParagraph(f"<b>{_escape_xml(ptext)}</b>", styles["ProposalBody"]))
        elif ptype == 'bullet':
            bullet_text = _markdown_bold_to_reportlab(ptext)
            story.append(RLParagraph(f"\u2022  {bullet_text}", styles["ProposalBullet"]))
        else:
            formatted = _markdown_bold_to_reportlab(ptext)
            if formatted.strip():
                story.append(RLParagraph(formatted, styles["ProposalBody"]))

        i += 1


def _add_pdf_table(story, headers, rows, styles):
    """Add a styled table to PDF story."""
    # Build table data with styled paragraphs
    header_row = [RLParagraph(f"<b>{_escape_xml(h)}</b>", styles["TableHeader"]) for h in headers]
    data = [header_row]
    for row in rows:
        data.append([RLParagraph(_escape_xml(str(c)), styles["TableCell"]) for c in row])

    num_cols = len(headers)
    available_width = 6.5 * inch
    col_width = available_width / num_cols

    table = RLTable(data, colWidths=[col_width] * num_cols)
    table_style = [
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1B2A4A")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("ALIGN", (0, 0), (-1, 0), "CENTER"),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#D1D5DB")),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
    ]
    # Alternating row colors
    for i in range(1, len(data)):
        if i % 2 == 0:
            table_style.append(("BACKGROUND", (0, i), (-1, i), colors.HexColor("#F3F4F6")))

    table.setStyle(TableStyle(table_style))
    story.append(table)


# ============================================================
# Utility functions
# ============================================================

def _escape_xml(text: str) -> str:
    """Escape special XML characters for ReportLab Paragraph."""
    text = text.replace("&", "&amp;")
    text = text.replace("<", "&lt;")
    text = text.replace(">", "&gt;")
    text = text.replace('"', "&quot;")
    return text


def _strip_markdown_bold(text: str) -> str:
    """Remove markdown bold markers (**text**) for DOCX output."""
    return re.sub(r"\*\*(.+?)\*\*", r"\1", text)


def _markdown_bold_to_reportlab(text: str) -> str:
    """Convert markdown bold (**text**) to ReportLab <b>text</b> tags."""
    text = _escape_xml(text)
    text = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", text)
    return text


def _html_to_paragraphs(html_content: str) -> list:
    """Convert HTML content into structured paragraph dicts."""
    if not html_content:
        return []

    paragraphs = []
    text = re.sub(r'<br\s*/?>', '\n', html_content)
    text = re.sub(r'<(?:strong|b)>(.*?)</(?:strong|b)>', r'**\1**', text, flags=re.DOTALL)
    text = re.sub(r'<(?:em|i)>(.*?)</(?:em|i)>', r'_\1_', text, flags=re.DOTALL)
    text = re.sub(r'<li>(.*?)</li>', r'\n- \1\n', text, flags=re.DOTALL)

    for level in range(1, 4):
        text = re.sub(
            rf'<h{level}[^>]*>(.*?)</h{level}>',
            lambda m, l=level: f'\n{"#" * l} {m.group(1).strip()}\n',
            text, flags=re.DOTALL,
        )

    text = re.sub(r'<[^>]+>', '', text)
    text = html_unescape(text)
    text = re.sub(r'\n{3,}', '\n\n', text)

    for line in text.split('\n'):
        line = line.strip()
        if not line:
            paragraphs.append({'type': 'spacer', 'text': ''})
        elif line.startswith('### '):
            paragraphs.append({'type': 'h3', 'text': line[4:].strip()})
        elif line.startswith('## '):
            paragraphs.append({'type': 'h2', 'text': line[3:].strip()})
        elif line.startswith('# '):
            paragraphs.append({'type': 'h1', 'text': line[2:].strip()})
        elif line.startswith('- ') or line.startswith('* ') or line.startswith('\u2022 '):
            prefix_len = 2
            if line.startswith('\u2022 '):
                prefix_len = 2
            paragraphs.append({'type': 'bullet', 'text': line[prefix_len:].strip()})
        elif re.match(r'^\d+\.\s', line):
            paragraphs.append({'type': 'numbered', 'text': re.sub(r'^\d+\.\s*', '', line).strip()})
        else:
            paragraphs.append({'type': 'text', 'text': line})

    return paragraphs
